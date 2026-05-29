"""Audit Booklet Generator - produces comprehensive .docx audit booklet with firm stamp."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
from pathlib import Path


def _set_shading(cell, color_hex):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(255,255,255)
        _set_shading(c, "1B3A5C")
    for ri, rd in enumerate(rows):
        for ci, val in enumerate(rd):
            c = t.rows[ri+1].cells[ci]
            c.text = str(val) if val else ""
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
            if ri % 2 == 1: _set_shading(c, "F0F4F8")
    return t


def _add_firm_header(doc, firm_name, firm_reg_no):
    """Add a firm identity header at the top of the document."""
    if not firm_name:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(firm_name)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(27, 58, 92)
    if firm_reg_no:
        run2 = p.add_run(f"  |  Reg. No.: {firm_reg_no}")
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(100, 100, 100)
    # Thin separator line
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C8965A')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_firm_footer_paragraph(doc, firm_name, firm_reg_no, generated_str):
    """Add a small footer note at the very end."""
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parts = []
    if firm_name:
        parts.append(firm_name)
    if firm_reg_no:
        parts.append(f"Reg. No. {firm_reg_no}")
    parts.append(f"Generated: {generated_str}")
    run = p.add_run("  ·  ".join(parts))
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(140, 140, 140)
    run.italic = True


def generate_booklet(eng, tasks, comments_by_task, reviews_by_task, queries, team, path,
                     firm_name="", firm_reg_no=""):
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.0)
    style = doc.styles["Normal"]; style.font.name = "Arial"; style.font.size = Pt(10)

    generated_str = datetime.now().strftime('%d-%m-%Y %H:%M')

    # Firm header on cover
    _add_firm_header(doc, firm_name, firm_reg_no)

    # Cover
    for _ in range(5): doc.add_paragraph("")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AUDIT BOOKLET"); r.bold = True; r.font.size = Pt(28); r.font.color.rgb = RGBColor(27,58,92)
    doc.add_paragraph("")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(eng.get("title","")); r.font.size = Pt(16); r.font.color.rgb = RGBColor(80,80,80)
    for txt in [f"Client: {eng.get('client_name','N/A')}", f"Financial Year: {eng.get('financial_year','N/A')}", f"Generated: {generated_str}"]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run(txt).font.size = Pt(11)

    # Firm stamp below cover details
    if firm_name:
        doc.add_paragraph("")
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        stamp = firm_name
        if firm_reg_no:
            stamp += f"  |  Reg. No.: {firm_reg_no}"
        run = p.add_run(stamp)
        run.font.size = Pt(9); run.italic = True; run.font.color.rgb = RGBColor(100,100,100)

    doc.add_page_break()

    # TOC
    doc.add_heading("Table of Contents", level=1)
    for item in ["1. Engagement Details","2. Audit Team","3. Task Execution Summary","4. Detailed Task Log","5. Comments & Review Trail","6. Query Sheet","7. Working Papers Index"]:
        doc.add_paragraph(item)
    doc.add_page_break()

    # 1. Engagement Details
    doc.add_heading("1. Engagement Details", level=1)
    _table(doc, ["Field","Details"], [
        ("Client Name", eng.get("client_name","")), ("Title", eng.get("title","")),
        ("Type", eng.get("engagement_type","")), ("Financial Year", eng.get("financial_year","")),
        ("Period", f"{eng.get('period_from','N/A')} to {eng.get('period_to','N/A')}"),
        ("Status", eng.get("status","")), ("PAN", eng.get("pan","")), ("GSTIN", eng.get("gstin","")),
        ("Prepared by", firm_name), ("Firm Reg. No.", firm_reg_no),
    ])
    doc.add_page_break()

    # 2. Team
    doc.add_heading("2. Audit Team", level=1)
    _table(doc, ["Sr.","Name","Role","Email"],
           [(str(i+1), m.get("full_name",""), m.get("role",""), m.get("email",""))
            for i, m in enumerate(team)])
    doc.add_page_break()

    # 3. Summary
    doc.add_heading("3. Task Execution Summary", level=1)
    sc = {}
    for t in tasks: sc[t.get("status","Unknown")] = sc.get(t.get("status","Unknown"),0)+1
    _table(doc, ["Status","Count"], [(s,str(c)) for s,c in sc.items()] + [("Total",str(len(tasks)))])
    doc.add_page_break()

    # 4. Task Log
    doc.add_heading("4. Detailed Task Log", level=1)
    _table(doc, ["ID","Area","Task","Assigned To","Status","Priority","WP Ref"],
        [(str(t.get("id","")), t.get("area",""), t.get("title",""),
          t.get("assignee_name","Unassigned"),
          t.get("status",""), t.get("priority",""), t.get("working_paper_ref",""))
         for t in tasks])
    doc.add_page_break()

    # 5. Comments & Reviews
    doc.add_heading("5. Comments & Review Trail", level=1)
    for t in tasks:
        tid = t.get("id")
        tc = comments_by_task.get(tid, []); tr = reviews_by_task.get(tid, [])
        if not tc and not tr: continue
        doc.add_heading(f"Task #{tid}: {t.get('title','')}", level=2)
        if tc:
            doc.add_heading("Comments", level=3)
            _table(doc, ["Author","Comment","Query?","Date"],
                [(c.get("author_name",""), c.get("content",""), "Yes" if c.get("is_query") else "", c.get("created_at","")) for c in tc])
        if tr:
            doc.add_heading("Reviews", level=3)
            _table(doc, ["Reviewer","Action","Remarks","Date"],
                [(r.get("reviewer_name",""), r.get("action",""), r.get("remarks",""), r.get("reviewed_at","")) for r in tr])
    doc.add_page_break()

    # 6. Query Sheet
    doc.add_heading("6. Query Sheet", level=1)
    if queries:
        _table(doc, ["Sr.","Query","Response","Status","Raised By","Date","Task Ref"],
            [(str(q.get("sr_no","")), q.get("query_text",""), q.get("response","") or "Pending", q.get("status",""), q.get("raised_by_name",""), q.get("raised_date",""), q.get("task_reference","")) for q in queries])
    else:
        doc.add_paragraph("No queries raised for this engagement.")
    doc.add_page_break()

    # 7. WP Index
    doc.add_heading("7. Working Papers Index", level=1)
    wps = [(str(i+1), t.get("working_paper_ref",""), t.get("title",""), t.get("area",""), t.get("status","")) for i,t in enumerate(tasks) if t.get("working_paper_ref")]
    if wps:
        _table(doc, ["Sr.","WP Reference","Task","Area","Status"], wps)
    else:
        doc.add_paragraph("No working paper references recorded.")

    # Firm footer on last page
    _add_firm_footer_paragraph(doc, firm_name, firm_reg_no, generated_str)

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path
